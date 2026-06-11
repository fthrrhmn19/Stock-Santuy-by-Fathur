from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/Panduan_Stock_Santuy_Analysis.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(22, 34, 52)
MUTED = RGBColor(96, 112, 132)
GREEN = RGBColor(34, 139, 92)
RED = RGBColor(174, 44, 58)
GOLD = RGBColor(122, 90, 0)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
PALE_GREEN = "EAF6EF"
PALE_RED = "FCEDEF"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="C8D2E0"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, before=0, after=6, line=1.25, align=None, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    set_paragraph_spacing(p, before, after, line)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, after=4, line=1.25)
        run = p.add_run(item)
        set_run_font(run, size=11, color=INK)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_spacing(p, after=4, line=1.25)
        run = p.add_run(item)
        set_run_font(run, size=11, color=INK)


def add_callout(doc, title, body, fill=LIGHT_GRAY, accent_color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="D2DAE6")
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=2, line=1.25)
    r = p.add_run(title)
    set_run_font(r, size=11, color=accent_color, bold=True)
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, after=0, line=1.25)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    add_para(doc, "", after=3)
    return table


def cell_text(cell, text, bold=False, color=INK, size=10.5, align=None):
    p = cell.paragraphs[0]
    p.text = ""
    set_paragraph_spacing(p, after=0, line=1.18)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_matrix(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], header_fill)
        cell_text(table.rows[0].cells[i], header, bold=True, color=DARK_BLUE, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, size=10.2)
    add_para(doc, "", after=4)
    return table


def add_label_detail(doc, rows, label_width=2700, detail_width=6660):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [label_width, detail_width])
    set_table_borders(table)
    for label, detail in rows:
        cells = table.add_row().cells
        if len(table.rows) == 1:
            mark_header_row(table.rows[0])
        shade_cell(cells[0], LIGHT_BLUE)
        cell_text(cells[0], label, bold=True, color=DARK_BLUE)
        cell_text(cells[1], detail)
    add_para(doc, "", after=4)
    return table


def add_section_break(doc):
    doc.add_page_break()


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
        sec.header_distance = Inches(0.492)
        sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def set_header_footer(doc):
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0]
        hp.text = ""
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(hp, after=0, line=1.0)
        run = hp.add_run("Stock Santuy Analysis | Buku Panduan")
        set_run_font(run, size=9, color=MUTED, bold=True)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(fp, after=0, line=1.0)
        r1 = fp.add_run("Halaman ")
        set_run_font(r1, size=9, color=MUTED)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "1"
        run.append(text)
        fld.append(run)
        fp._p.append(fld)


def build_doc():
    doc = Document()
    style_document(doc)

    # Cover
    add_para(doc, "BUKU PANDUAN", size=10.5, color=GOLD, bold=True, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Stock Santuy Analysis")
    add_para(doc, "Panduan pengguna, latihan belajar, dan cara membaca sinyal web", size=14, color=MUTED, after=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Disusun untuk membantu pengguna memahami dashboard analisis saham Indonesia berbasis OHLCV, screener, market mover, chart, news, risk scoring, dan position sizing.", size=11, color=INK, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_matrix(
        doc,
        ["Item", "Keterangan"],
        [
            ("Nama web", "Stock Santuy Analysis by Fathur"),
            ("Versi panduan", "8 Juni 2026"),
            ("Fokus belajar", "Day trade, swing, long term, potential bagger, dan manajemen risiko"),
            ("Catatan data", "Data gratis bisa delayed dan bukan pengganti orderbook broker atau data IDX resmi"),
        ],
        [2400, 6960],
        header_fill=LIGHT_BLUE,
    )
    add_callout(
        doc,
        "Prinsip utama",
        "Gunakan web ini sebagai alat bantu edukasi dan riset. Keputusan transaksi tetap harus divalidasi dengan orderbook, berita resmi, laporan keuangan, dan risk management pribadi.",
        fill=PALE_GREEN,
        accent_color=GREEN,
    )
    add_section_break(doc)

    add_heading(doc, "Daftar Isi", 1)
    add_matrix(
        doc,
        ["Bagian", "Isi utama"],
        [
            ("1", "Quick start: cara mulai dalam 15 menit"),
            ("2", "Peta fitur web dan fungsi tiap panel"),
            ("3", "Cara menjalankan analisa saham"),
            ("4", "Cara membaca chart, mode, dan indikator"),
            ("5", "Screener, market mover, IHSG, news, dan alert"),
            ("6", "Trade setup, position sizing, dan checklist risiko"),
            ("7", "Latihan belajar 7 hari"),
            ("8", "Troubleshooting, batas akurasi, dan glosarium"),
        ],
        [1600, 7760],
    )

    add_heading(doc, "Cara Membaca Panduan Ini", 1)
    add_bullets(doc, [
        "Kalau baru pertama pakai, mulai dari Quick Start lalu praktik dengan ticker besar seperti BBCA, BBRI, BMRI, TLKM, atau ASII.",
        "Kalau fokus trading harian, baca bagian IHSG, Live 5M, Trading Screener, Top Gainer/Loser, dan Position Size.",
        "Kalau fokus swing atau long term, baca bagian mode Swing, Long Term, Potential Bagger, valuasi, fundamental, dan news ekspansi.",
        "Setiap sinyal dipakai sebagai filter awal. Jangan langsung transaksi hanya karena skor terlihat tinggi.",
    ])
    add_callout(
        doc,
        "Rumus belajar yang aman",
        "Lihat konteks pasar dulu, pilih saham dari screener, buka analisa detail, baca chart dan mode, tentukan entry/stop/target, lalu hitung ukuran posisi. Urutan ini menjaga kamu tidak mengejar saham hanya karena sedang ramai.",
        fill=LIGHT_GRAY,
    )
    add_section_break(doc)

    add_heading(doc, "1. Quick Start 15 Menit", 1)
    add_para(doc, "Bagian ini adalah jalur tercepat untuk memahami cara memakai web dari nol sampai bisa membaca satu analisa saham secara utuh.")
    add_numbers(doc, [
        "Buka dashboard. Jika login aktif, masukkan password akses yang sudah disiapkan.",
        "Lihat status provider dan market clock di bar status. Pastikan tidak ada pesan error koneksi.",
        "Masukkan ticker di Quick Analyze, misalnya BBCA, lalu tekan Analisa.",
        "Baca ringkasan harga, perubahan, label skor, dan Final Verdict.",
        "Buka tab mode Day Trade, Swing, Long Term, dan Potential Bagger untuk melihat mode yang paling cocok.",
        "Cek Grafik Harga. Gunakan Live untuk candle 5 menit, 1D untuk harian, 1W untuk mingguan, dan 1Bln sampai 1Th untuk konteks panjang.",
        "Cek Trade Setup Analyzer: entry, breakout, stop loss, target, dan invalidasi.",
        "Hitung Position Size sesuai modal dan risiko per transaksi.",
    ])
    add_matrix(
        doc,
        ["Kalau terlihat", "Artinya", "Tindakan belajar"],
        [
            ("Skor tinggi", "Setup teknikal relatif kuat menurut aturan web", "Tetap cek trend, volume, IHSG, dan risk level"),
            ("Volume spike", "Aktivitas transaksi di atas rata-rata", "Cek apakah harga ikut valid atau hanya noise"),
            ("Downtrend", "Struktur harga masih lemah", "Jangan buru-buru buy; tunggu konfirmasi reversal"),
            ("Data delayed", "Provider gratis tidak realtime resmi", "Bandingkan dengan broker sebelum eksekusi"),
        ],
        [2200, 3600, 3560],
    )

    add_heading(doc, "2. Peta Fitur Web", 1)
    add_para(doc, "Dashboard dibagi menjadi beberapa area. Setiap area punya tugas berbeda, jadi gunakan sesuai kebutuhan.")
    add_label_detail(doc, [
        ("Navbar", "Akses cepat ke Cari, Analisa, Screener, IHSG, News, dan Alert."),
        ("Quick Analyze", "Tempat memasukkan ticker untuk membuka analisa lengkap."),
        ("Trading Screener", "Daftar kandidat momentum, breakout, volume spike, dan swing jangka pendek."),
        ("Investment Screener", "Daftar kandidat trend panjang, stabilitas, dan likuiditas."),
        ("Potential Bagger", "Filter awal saham dengan potensi ekspansi; tetap wajib validasi lapkeu dan berita resmi."),
        ("AI Market Scanner", "Top Gainer, Top Loser, Top Value, dan Top Volume berdasarkan market mover."),
        ("IHSG", "Konteks arah pasar. Saat IHSG lemah, seleksi entry perlu lebih ketat."),
        ("Full Analisa Saham", "Panel detail untuk membaca harga, chart, indikator, risiko, valuasi, setup, dan strategi."),
        ("Market News", "Berita dan sinyal ekspansi/lapkeu untuk validasi katalis."),
        ("Auto Email Alert", "Pengecekan kandidat alert jika environment email sudah diisi."),
    ])

    add_heading(doc, "3. Status Provider, Market Clock, dan Auto Refresh", 1)
    add_para(doc, "Bar status di bawah header memberi tahu apakah provider aktif, fase market IDX, status auto refresh, dan tombol refresh scanner.")
    add_bullets(doc, [
        "Provider aktif berarti endpoint data dapat diakses.",
        "Market clock mengikuti timezone Asia/Jakarta dan sesi reguler IDX.",
        "Auto ON berarti web polling data berkala. Versi terbaru tetap mengecek data delayed walaupun market sedang istirahat atau tutup.",
        "Refresh Scanner digunakan jika ingin memaksa update daftar screener saat itu juga.",
        "Jika data tidak berubah, penyebabnya bisa karena candle 5 menit belum terbentuk, provider delayed, atau cache data belum lewat.",
    ])
    add_callout(
        doc,
        "Live bukan tick-by-tick broker",
        "Chart Live memakai candle intraday 5 menit saat tersedia. Ini membantu pemantauan, tetapi bukan streaming tick-by-tick resmi IDX. Untuk eksekusi, tetap cek broker atau data resmi.",
        fill=PALE_RED,
        accent_color=RED,
    )
    add_section_break(doc)

    add_heading(doc, "4. Cara Menjalankan Analisa Saham", 1)
    add_para(doc, "Gunakan alur berikut setiap kali menganalisa ticker baru.")
    add_numbers(doc, [
        "Tentukan tujuan: day trade, swing, long term, atau hanya belajar membaca struktur.",
        "Masukkan ticker di Quick Analyze atau klik tombol Analisa dari kartu screener.",
        "Baca judul analisa untuk memastikan ticker dan nama emiten benar.",
        "Cek harga terakhir, perubahan, open, high, low, dan data status.",
        "Baca Insight Utama: trend, mode terbaik, faktor dominan, invalidasi utama.",
        "Cocokkan Final Verdict dengan chart dan risk panel.",
    ])
    add_matrix(
        doc,
        ["Mode", "Cocok untuk", "Yang wajib dicek"],
        [
            ("Day Trade", "Momentum cepat intraday", "Live 5M, IHSG, volume, support intraday, stop ketat"),
            ("Swing", "Setup beberapa hari sampai minggu", "1D, EMA, RSI, MACD, breakout, pullback, target"),
            ("Long Term", "Akumulasi bertahap", "Trend panjang, EMA50/200, valuasi, fundamental, news"),
            ("Potential Bagger", "Filter awal saham ekspansi", "Stage, trend, likuiditas, katalis bisnis, laporan keuangan"),
        ],
        [1900, 3300, 4160],
    )

    add_heading(doc, "5. Membaca Chart Harga", 1)
    add_para(doc, "Chart adalah tempat memvalidasi sinyal. Jangan hanya membaca skor; lihat apakah struktur harga mendukung.")
    add_label_detail(doc, [
        ("Live", "Candle intraday 5 menit. Cocok untuk melihat ritme harian saat data tersedia."),
        ("1D", "Candle harian. Cocok untuk swing, support-resistance, dan tren beberapa minggu."),
        ("1W", "Candle mingguan. Cocok untuk melihat trend besar dan area akumulasi."),
        ("1Bln", "Candle bulanan. Cocok untuk investor dan konteks long term."),
        ("3Bln/6Bln/1Th", "Agregasi candle panjang. Cocok untuk melihat siklus besar, bukan entry cepat."),
    ])
    add_bullets(doc, [
        "Candle hijau menunjukkan close lebih tinggi dari open; candle merah menunjukkan close lebih rendah dari open.",
        "Volume bar membantu melihat apakah pergerakan harga didukung transaksi besar.",
        "EMA20 di chart membantu membaca arah pendek sampai menengah.",
        "Marker stop loss dan target membantu menilai risk-reward secara visual.",
        "Jika candle melebar jauh dari support, hindari mengejar harga tanpa rencana stop.",
    ])

    add_heading(doc, "6. Membaca Indikator Utama", 1)
    add_matrix(
        doc,
        ["Indikator", "Fungsi", "Cara pakai praktis"],
        [
            ("Trend", "Membaca arah struktur harga", "Prioritaskan buy saat uptrend atau mulai reversal valid"),
            ("RSI", "Mengukur momentum dan area overbought/oversold", "RSI tinggi tidak selalu jual; cek trend dan volume"),
            ("MACD", "Membaca momentum lanjutan", "Histogram membaik membantu konfirmasi trend"),
            ("Relative Volume", "Membandingkan volume terbaru vs rata-rata", "RVOL tinggi + candle kuat lebih valid daripada volume tinggi tapi harga lemah"),
            ("Support", "Area harga yang sering ditahan buyer", "Entry pullback biasanya dicari dekat support sehat"),
            ("Resistance", "Area harga yang sering ditahan seller", "Breakout perlu volume dan close kuat"),
            ("ATR/Volatilitas", "Mengukur lebar gerakan normal", "Dipakai untuk stop loss dan ukuran posisi"),
        ],
        [1800, 2920, 4640],
    )
    add_callout(
        doc,
        "Kunci membaca volume",
        "Volume besar harus dibandingkan dengan hasil harga. Kalau volume besar tetapi candle sempit dan close lemah, itu bisa menjadi tanda supply atau distribusi, bukan otomatis bullish.",
        fill=LIGHT_GRAY,
    )
    add_section_break(doc)

    add_heading(doc, "7. Membaca Screener", 1)
    add_para(doc, "Screener adalah pintu masuk untuk menemukan kandidat. Screener bukan perintah beli; anggap sebagai daftar saham yang layak dibuka analisa detailnya.")
    add_matrix(
        doc,
        ["Screener", "Tujuan", "Kapan dipakai"],
        [
            ("Trading Ideas", "Mencari kandidat momentum, breakout, volume spike, dan swing", "Saat ingin mencari saham aktif untuk hari ini atau beberapa hari ke depan"),
            ("Investment Ideas", "Mencari saham dengan trend panjang, likuiditas, dan stabilitas", "Saat ingin membuat watchlist akumulasi"),
            ("Potential Bagger", "Mencari kandidat ekspansi dan trend kuat", "Saat ingin riset lebih dalam ke laporan keuangan dan katalis"),
        ],
        [2400, 3760, 3200],
    )
    add_bullets(doc, [
        "Klik Analisa pada kartu saham untuk membuka detail.",
        "Bandingkan minimal 3 kandidat sebelum memilih satu yang paling bersih.",
        "Jangan pilih hanya karena score paling tinggi; risk-reward dan likuiditas tetap penting.",
        "Untuk saham kecil, selalu cek spread bid-offer di broker.",
    ])

    add_heading(doc, "8. Membaca AI Market Scanner", 1)
    add_matrix(
        doc,
        ["Kategori", "Definisi praktis", "Cara menggunakannya"],
        [
            ("Top Gainer", "Saham dengan kenaikan persentase terbesar", "Cari momentum, tetapi waspada saham yang sudah terlalu jauh"),
            ("Top Loser", "Saham dengan penurunan persentase terbesar", "Cari risiko, panic selling, atau kandidat rebound hanya jika ada konfirmasi"),
            ("Top Value", "Saham dengan nilai transaksi terbesar", "Cari saham yang menjadi pusat uang besar hari itu"),
            ("Top Volume", "Saham dengan volume/lot transaksi terbesar", "Cari aktivitas ramai; cek apakah harganya ikut valid"),
        ],
        [2000, 3400, 3960],
    )
    add_callout(
        doc,
        "Market mover perlu konteks",
        "Top Gainer yang naik besar belum tentu aman. Top Value yang besar belum tentu bullish. Selalu gabungkan dengan chart, volume, IHSG, dan rencana stop loss.",
        fill=PALE_RED,
        accent_color=RED,
    )

    add_heading(doc, "9. Membaca IHSG", 1)
    add_para(doc, "IHSG membantu membaca angin pasar. Saat IHSG kuat, breakout saham lebih mudah lanjut. Saat IHSG lemah, banyak setup bagus bisa gagal karena tekanan market.")
    add_bullets(doc, [
        "Jika IHSG hijau dan membentuk higher high/higher low, setup long lebih nyaman.",
        "Jika IHSG merah tajam, kurangi agresivitas entry dan perkecil posisi.",
        "Jika IHSG sideways, pilih saham dengan katalis atau volume yang lebih jelas.",
        "Jika IHSG turun tetapi saham tetap kuat, saham itu layak masuk watchlist relative strength.",
    ])
    add_section_break(doc)

    add_heading(doc, "10. Trade Setup Analyzer", 1)
    add_para(doc, "Panel ini menerjemahkan analisa menjadi area tindakan. Gunakan sebagai peta, bukan angka sakral.")
    add_matrix(
        doc,
        ["Field", "Makna", "Cara pakai"],
        [
            ("Buy on weakness", "Area beli saat harga pullback", "Lebih santuy daripada mengejar breakout"),
            ("Breakout entry", "Area masuk agresif saat tembus resistance", "Butuh volume dan candle close kuat"),
            ("Cut loss", "Batas invalidasi setup", "Wajib dihormati jika rencana gagal"),
            ("TP1/TP2/TP3", "Target bertahap", "Bisa dipakai untuk jual sebagian atau trailing stop"),
            ("Invalidasi", "Level yang membatalkan skenario", "Jika tembus, evaluasi ulang tanpa emosi"),
        ],
        [1900, 3380, 4080],
    )

    add_heading(doc, "11. Position Size", 1)
    add_para(doc, "Position sizing membantu menjaga risiko tetap kecil walaupun analisa salah.")
    add_numbers(doc, [
        "Isi modal total yang siap dipakai.",
        "Isi risiko per transaksi, misalnya 0,5% sampai 2%.",
        "Klik Hitung.",
        "Web menghitung estimasi lot berdasarkan jarak entry dan stop loss.",
        "Jika lot terlalu besar atau biaya terlalu tinggi, turunkan risiko atau tunggu entry lebih dekat support.",
    ])
    add_matrix(
        doc,
        ["Contoh keputusan", "Lebih baik dilakukan"],
        [
            ("Stop loss terlalu jauh", "Kurangi lot atau tunggu harga pullback"),
            ("Risk-reward kurang menarik", "Cari setup lain; jangan memaksa entry"),
            ("Saham tidak likuid", "Perkecil posisi dan cek bid-offer"),
            ("Market sedang lemah", "Turunkan risiko per posisi"),
        ],
        [3300, 6060],
    )

    add_heading(doc, "12. Valuasi dan Fundamental", 1)
    add_para(doc, "Panel valuasi dan fundamental adalah validasi awal. Data rasio publik membantu screening, tetapi final tetap harus ke laporan keuangan resmi.")
    add_label_detail(doc, [
        ("Fair value teknikal", "Proxy teknikal, bukan intrinsic value ala Graham/Klarman."),
        ("EPS, PER, PBV", "Rasio dasar untuk membaca harga terhadap laba dan buku."),
        ("ROE", "Mengukur profitabilitas terhadap ekuitas."),
        ("DER", "Membaca leverage atau beban utang/liabilitas terhadap ekuitas."),
        ("Trading value", "Nilai transaksi; membantu menilai likuiditas harian."),
        ("Data confidence", "Kualitas kelengkapan data yang tersedia dari sumber publik."),
    ])
    add_callout(
        doc,
        "Aturan aman untuk fundamental",
        "Jika ingin investasi atau potential bagger, jangan berhenti di panel web. Baca laporan keuangan, public expose, aksi korporasi, dan berita resmi emiten.",
        fill=PALE_GREEN,
        accent_color=GREEN,
    )
    add_section_break(doc)

    add_heading(doc, "13. E-book Strategy Engine", 1)
    add_para(doc, "Panel ini menampilkan sinyal yang bisa dihitung dari OHLCV dan aturan strategi yang sudah dipetakan ke web.")
    add_matrix(
        doc,
        ["Sinyal", "Makna ringkas", "Yang perlu diingat"],
        [
            ("VPA bullish validation", "Candle naik, spread kuat, volume di atas rata-rata, close kuat", "Valid jika harga dan volume searah"),
            ("Breakout volume confirmation", "Harga menembus resistance dengan volume", "Lebih kuat jika close tidak balik ke bawah resistance"),
            ("No demand warning", "Kenaikan kecil pada volume rendah", "Waspada demand lemah"),
            ("Stopping volume watch", "Tekanan turun mulai diserap volume besar", "Butuh konfirmasi candle berikutnya"),
            ("Weinstein Stage proxy", "Membaca stage trend memakai SMA150 harian", "Stage 2 lebih menarik, Stage 4 dihindari"),
            ("Minervini Trend Template", "Checklist trend kuat jangka menengah", "RS ranking belum lengkap"),
            ("Boxer price-volume surge", "Momentum cepat dengan harga dan volume meningkat", "Cocok untuk watchlist aktif"),
        ],
        [2500, 3600, 3260],
    )

    add_heading(doc, "14. News, Daily Snapshot, dan Alert", 1)
    add_label_detail(doc, [
        ("Market News", "Berita dipakai untuk membantu membaca katalis, terutama ekspansi, laporan keuangan, atau aksi korporasi."),
        ("Sinyal Ekspansi/Lapkeu", "Area khusus untuk kata kunci yang relevan dengan potential bagger dan long term."),
        ("Daily Snapshot", "Tombol Harian membuat PNG rekomendasi berdasarkan scanner terbaru."),
        ("Auto Email Alert", "Jika konfigurasi email aktif, sistem bisa mengirim alert terjadwal."),
    ])
    add_bullets(doc, [
        "Berita positif harus dikonfirmasi oleh volume dan struktur harga.",
        "Berita tanpa volume sering tidak cukup untuk entry.",
        "Alert adalah pengingat untuk cek chart, bukan perintah transaksi.",
        "Daily Snapshot cocok untuk jurnal atau catatan watchlist harian.",
    ])

    add_section_break(doc)
    add_heading(doc, "15. Rutinitas Harian yang Disarankan", 1)
    add_matrix(
        doc,
        ["Waktu", "Yang dicek", "Tujuan"],
        [
            ("Sebelum market", "News, IHSG terakhir, watchlist, level support-resistance", "Menyiapkan rencana, bukan bereaksi panik"),
            ("Sesi I", "Top Value, Top Volume, Trading Ideas, Live 5M", "Mencari saham yang benar-benar aktif"),
            ("Istirahat", "Review saham yang bertahan kuat dan volume masuk", "Memilih kandidat sesi II"),
            ("Sesi II", "Breakout, pullback, risk-reward, position size", "Eksekusi hanya jika setup lengkap"),
            ("Setelah market", "Daily chart, jurnal, evaluasi stop/target", "Belajar dari keputusan hari itu"),
        ],
        [1700, 4300, 3360],
    )
    add_callout(
        doc,
        "Kebiasaan terbaik",
        "Sebelum klik buy, tulis dulu alasan entry, level invalidasi, target, ukuran posisi, dan kondisi yang membuat kamu batal entry.",
        fill=LIGHT_GRAY,
    )

    add_heading(doc, "16. Latihan Belajar 7 Hari", 1)
    add_matrix(
        doc,
        ["Hari", "Latihan", "Output"],
        [
            ("1", "Buka BBCA, BBRI, BMRI. Bandingkan trend, RSI, volume, dan Final Verdict.", "Catatan 3 perbedaan utama"),
            ("2", "Bandingkan IHSG dengan 3 saham pilihan.", "Saham yang lebih kuat/lemah dari market"),
            ("3", "Ambil 5 saham dari Trading Ideas lalu buka analisa satu per satu.", "Ranking pribadi berdasarkan risk-reward"),
            ("4", "Cek Top Gainer, Top Loser, Top Value, Top Volume.", "Pahami bedanya naik besar vs transaksi besar"),
            ("5", "Latih Position Size dengan modal simulasi Rp 10 juta dan risiko 1%.", "Lot maksimal untuk 3 saham"),
            ("6", "Cari satu kandidat long term dan cek valuasi/news/fundamental.", "Daftar data tambahan yang harus dicari"),
            ("7", "Buat watchlist mingguan dari 5 saham terbaik.", "Rencana entry, stop, target, dan alasan batal"),
        ],
        [900, 5860, 2600],
    )
    add_section_break(doc)

    add_heading(doc, "17. Checklist Sebelum Entry", 1)
    add_matrix(
        doc,
        ["Cek", "Pertanyaan"],
        [
            ("Market", "Apakah IHSG mendukung atau sedang menekan mayoritas saham?"),
            ("Trend", "Apakah saham uptrend, sideways sehat, atau masih downtrend?"),
            ("Volume", "Apakah volume masuk mendukung arah harga?"),
            ("Entry", "Apakah entry dekat support/pullback atau breakout valid?"),
            ("Stop", "Apakah level cut loss jelas dan masuk akal?"),
            ("Risk", "Apakah ukuran posisi sesuai risiko maksimal?"),
            ("Katalis", "Apakah ada news/lapkeu/aksi korporasi yang perlu dicek?"),
            ("Likuiditas", "Apakah bid-offer dan volume cukup untuk keluar masuk posisi?"),
        ],
        [1700, 7660],
    )

    add_heading(doc, "18. Troubleshooting", 1)
    add_label_detail(doc, [
        ("Data tidak berubah", "Bisa karena candle 5 menit belum terbentuk, provider delayed, atau cache 60 detik belum lewat."),
        ("Chart kosong", "Cek koneksi, provider, ticker, dan apakah saham tersedia di Yahoo/Twelve Data."),
        ("Screener lama loading", "Scanner memanggil banyak ticker dan sumber publik; tunggu atau klik refresh lagi."),
        ("Login muncul", "Masukkan password akses jika login opsional diaktifkan."),
        ("Alert tidak terkirim", "Cek RESEND_API_KEY, email tujuan, email pengirim, jadwal Cloudflare, dan SITE_URL."),
        ("Angka beda dari broker", "Provider gratis tidak selalu realtime; broker/orderbook lebih dekat ke data eksekusi."),
    ])

    add_heading(doc, "19. Batas Akurasi dan Disclaimer", 1)
    add_bullets(doc, [
        "Data gratis bisa delayed, EOD, atau berbeda dari data broker.",
        "Dashboard bukan penasihat investasi dan bukan ajakan membeli atau menjual efek.",
        "Skor teknikal tidak menjamin harga akan naik.",
        "Potential bagger adalah filter awal, bukan kesimpulan bisnis final.",
        "Validasi final tetap membutuhkan laporan keuangan resmi, keterbukaan informasi IDX, prospek bisnis, dan kondisi market.",
    ])

    add_heading(doc, "20. Glosarium Singkat", 1)
    add_matrix(
        doc,
        ["Istilah", "Arti"],
        [
            ("OHLCV", "Open, High, Low, Close, Volume."),
            ("Support", "Area harga yang cenderung ditahan buyer."),
            ("Resistance", "Area harga yang cenderung ditahan seller."),
            ("Breakout", "Harga menembus resistance atau range penting."),
            ("Pullback", "Penurunan sementara dalam trend naik."),
            ("RVOL", "Relative volume, volume saat ini dibanding rata-rata."),
            ("ATR", "Average True Range, ukuran volatilitas."),
            ("Stop loss", "Level keluar saat skenario salah."),
            ("Risk-reward", "Perbandingan potensi rugi dan potensi untung."),
            ("Bagger", "Saham yang berpotensi naik berkali-kali lipat, perlu validasi fundamental kuat."),
        ],
        [2200, 7160],
    )
    add_callout(
        doc,
        "Penutup",
        "Web yang bagus membantu kamu melihat data lebih cepat. Trader atau investor yang baik tetap harus disiplin pada proses: validasi, rencana, risiko, dan evaluasi.",
        fill=PALE_GREEN,
        accent_color=GREEN,
    )

    set_header_footer(doc)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
